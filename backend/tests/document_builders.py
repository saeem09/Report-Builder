"""In-memory builders for real test documents.

Test inputs are generated with the same libraries that parse them, so the
suite is self-contained and needs no checked-in binary fixture files. The
PDF constants are minimal hand-written PDF literals, which avoids adding a
PDF-writing library that production code would never use.
"""

import io

from docx import Document

SINGLE_PAGE_PDF_BYTES = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 49 >>
stream
BT /F1 24 Tf 72 720 Td (Sprint review notes) Tj ET
endstream
endobj
trailer
<< /Size 6 /Root 1 0 R >>
%%EOF
"""

TWO_PAGE_PDF_BYTES = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R 6 0 R] /Count 2 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 49 >>
stream
BT /F1 24 Tf 72 720 Td (Sprint review notes) Tj ET
endstream
endobj
6 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 7 0 R >>
endobj
7 0 obj
<< /Length 42 >>
stream
BT /F1 24 Tf 72 720 Td (Action items) Tj ET
endstream
endobj
trailer
<< /Size 8 /Root 1 0 R >>
%%EOF
"""


NO_TEXT_PDF_BYTES = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 5 0 R >>
endobj
5 0 obj
<< /Length 0 >>
stream

endstream
endobj
trailer
<< /Size 6 /Root 1 0 R >>
%%EOF
"""


def build_docx_bytes(paragraphs, table_rows=()):
    """Build a .docx file in memory.

    paragraphs: an iterable of paragraph strings, added in order.
    table_rows: an iterable of equal-length row tuples appended as one table.
    """
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row_values in enumerate(table_rows):
            for cell_index, value in enumerate(row_values):
                table.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
