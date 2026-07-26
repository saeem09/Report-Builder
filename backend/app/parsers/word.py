import io
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .errors import DocumentParseError

CELL_SEPARATOR = "\t"
BLOCK_SEPARATOR = "\n"


def parse_docx(content: bytes) -> str:
    """Extract text from a Word document.

    Paragraphs come first in document order, then every table row rendered as
    tab-separated cells. python-docx does not include table cell text in
    document.paragraphs, and meeting documents routinely put action items in
    tables, so both are collected.
    """
    try:
        document = Document(io.BytesIO(content))
    except (BadZipFile, PackageNotFoundError, ValueError) as error:
        raise DocumentParseError(
            "Could not open the file as a Word (.docx) document."
        ) from error

    paragraph_lines = [paragraph.text for paragraph in document.paragraphs]
    table_lines = [
        CELL_SEPARATOR.join(cell.text for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    return BLOCK_SEPARATOR.join(paragraph_lines + table_lines)
